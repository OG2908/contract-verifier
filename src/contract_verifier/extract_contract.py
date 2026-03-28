"""Extract structured data from Hebrew purchase contract DOCX files."""
from __future__ import annotations

import re
import logging

from docx import Document

from .models import ContractData, CostLine, PaymentLine, parse_hebrew_amount

logger = logging.getLogger(__name__)


def extract(docx_path: str) -> ContractData:
    """Extract contract data from a Hebrew DOCX file."""
    doc = Document(docx_path)

    # Extract client details from Table 0 (parties table)
    client_name = _extract_client(doc)

    # Extract definitions from Table 2
    registration_fee, delivery_date = _extract_definitions(doc)

    # Extract Appendix A (property details) from table with "מספר דירה"
    appendix_a = _find_table_by_label(doc, "מספר דירה")
    if appendix_a is None:
        raise ExtractionError("Cannot find Appendix A table (looking for 'מספר דירה')")

    apartment_number = _get_row_text(appendix_a, "מספר דירה")
    floor = _get_row_text(appendix_a, "קומה")
    area_gross_sqm = _parse_area(_get_row_text(appendix_a, "שטח דירה"))
    balcony_sqm = _parse_area(_get_row_text(appendix_a, "שטח מרפסת"))
    total_purchase_price = _parse_price(_get_row_text(appendix_a, "סכום הרכישה"))
    total_costs_percentage = _extract_costs_percentage(appendix_a)

    # Extract Appendix D (payment schedule) from table with "פרויקט"
    appendix_d = _find_table_by_label(doc, "מחיר רכישה כולל")
    if appendix_d is None:
        raise ExtractionError("Cannot find Appendix D table (looking for 'מחיר רכישה כולל')")

    project_name = _extract_project_name(appendix_d)
    remaining_after_registration = _parse_price(
        _get_row_text(appendix_d, "נותר לשלם")
    )
    surcharge_percentage = _extract_surcharge_percentage(appendix_d)
    payment_lines = _extract_payment_lines(appendix_d)

    # Cross-check: Appendix D also has total price and registration
    d_total = _parse_price(_get_row_text(appendix_d, "מחיר רכישה כולל"))
    d_reg = _parse_price(_get_row_text(appendix_d, "דמי הרשמה"))

    logger.debug("Appendix A total: %s, Appendix D total: %s", total_purchase_price, d_total)
    logger.debug("Definitions reg fee: %s, Appendix D reg fee: %s", registration_fee, d_reg)

    return ContractData(
        client_name=client_name,
        apartment_number=apartment_number,
        floor=floor,
        area_gross_sqm=area_gross_sqm,
        balcony_sqm=balcony_sqm,
        total_purchase_price=total_purchase_price,
        total_costs_percentage=total_costs_percentage,
        cost_lines=[],  # Individual cost lines are not in the contract template
        registration_fee=registration_fee,
        remaining_after_registration=remaining_after_registration,
        surcharge_percentage=surcharge_percentage,
        payment_lines=payment_lines,
        project_name=project_name,
        delivery_date=delivery_date,
    )


def _find_table_by_label(doc: Document, label: str):
    """Find a table containing a specific Hebrew label."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if label in cell.text:
                    return table
    return None


def _get_row_text(table, label: str) -> str:
    """Find a row where any cell contains `label`, return value from the other cell."""
    for row in table.rows:
        cells = row.cells
        for i, cell in enumerate(cells):
            if label in cell.text.strip():
                # Return the first cell that has different content
                for j in range(len(cells)):
                    if j != i:
                        val = cells[j].text.strip()
                        if val and val != cell.text.strip():
                            return val
                # If all other cells are empty, check if value is embedded in label cell
                # e.g., "סכום העלויות הנלוות מתוך סכום הרכישה (8.5%):"
                return cell.text.strip()
    raise ExtractionError(f"Cannot find row with label '{label}' in table")



def _parse_area(text: str) -> float:
    """Parse area value like '37.07 מ"ר ברוטו' or '7.19 מ״ר'."""
    cleaned = re.sub(r'[\u200f\u200e\u200b\u00a0]', '', text)
    match = re.search(r'([\d,.]+)', cleaned)
    if not match:
        raise ExtractionError(f"Cannot parse area from: {text!r}")
    return float(match.group(1).replace(',', ''))


def _parse_price(text: str) -> float:
    """Parse a price value like '€122,224' or '2,000'."""
    cleaned = re.sub(r'[\u200f\u200e\u200b\u00a0]', '', text)
    # Remove everything except digits, commas, periods, euro sign
    match = re.search(r'€?([\d,]+)', cleaned)
    if not match:
        raise ExtractionError(f"Cannot parse price from: {text!r}")
    return parse_hebrew_amount(match.group(1))


def _extract_costs_percentage(table) -> float:
    """Extract the total costs percentage from Appendix A.

    Found in a row like: "סכום העלויות הנלוות מתוך סכום הרכישה (8.5%):"
    """
    for row in table.rows:
        for cell in row.cells:
            text = cell.text
            if "עלויות" in text:
                match = re.search(r'([\d.]+)%', text)
                if match:
                    return float(match.group(1))
    raise ExtractionError("Cannot find costs percentage in Appendix A")


def _extract_project_name(table) -> str:
    """Extract project name from Appendix D header row."""
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if "פרויקט" in cells:
            idx = cells.index("פרויקט")
            # Value is in the next cell
            if idx + 1 < len(cells) and cells[idx + 1]:
                return cells[idx + 1]
    raise ExtractionError("Cannot find project name in Appendix D")


def _extract_surcharge_percentage(table) -> float:
    """Extract surcharge percentage from Appendix D header.

    Found in text like: "בתוספת 2%"
    """
    for row in table.rows:
        for cell in row.cells:
            if "בתוספת" in cell.text:
                match = re.search(r'בתוספת\s+([\d.]+)%', cell.text)
                if match:
                    return float(match.group(1))
    raise ExtractionError("Cannot find surcharge percentage in Appendix D")


def _extract_payment_lines(table) -> list[PaymentLine]:
    """Extract payment schedule lines from Appendix D."""
    payment_labels = ["מקדמה", "תשלום ראשון", "תשלום שני", "תשלום שלישי"]
    payments = []

    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        full_row = " ".join(cells)

        for label in payment_labels:
            if label in full_row:
                # Extract percentage
                pct_match = re.search(r'(\d+)%', full_row)
                # Extract euro amounts
                amounts = re.findall(r'€([\d,]+)', full_row)

                if pct_match and len(amounts) >= 2:
                    payments.append(PaymentLine(
                        name=label,
                        percentage=float(pct_match.group(1)),
                        base_amount=parse_hebrew_amount(amounts[0]),
                        amount_with_surcharge=parse_hebrew_amount(amounts[1]),
                        notes=cells[-1] if len(cells) > 4 else "",
                    ))
                elif pct_match and len(amounts) == 1:
                    # Only one amount found
                    payments.append(PaymentLine(
                        name=label,
                        percentage=float(pct_match.group(1)),
                        base_amount=parse_hebrew_amount(amounts[0]),
                        amount_with_surcharge=0,
                        notes=cells[-1] if len(cells) > 4 else "",
                    ))
                break

    if not payments:
        raise ExtractionError("Cannot extract any payment lines from Appendix D")

    return payments


def _extract_client(doc: Document) -> str:
    """Extract client name from Table 0 (parties table).

    Row 0 contains buyer info in cell[1]:
    'נילי שטרן ביבר ת.ז. 60637642\nכתובת: ...\nדוא"ל: ...\nנייד: ...'
    """
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            for cell_text in cells:
                if "ת.ז" in cell_text or "ת.ז." in cell_text:
                    lines = cell_text.split("\n")
                    first_line = lines[0].strip()
                    id_match = re.search(r'ת\.?ז\.?\s*([\d]+)', first_line)
                    if id_match:
                        name = first_line[:id_match.start()].strip()
                        return name

    raise ExtractionError("Cannot extract client name from contract")


def _extract_definitions(doc: Document) -> tuple[float, str]:
    """Extract registration fee and delivery date from definitions table (Table 2).

    Registration fee: row with "דמי הקמה/ הרשמה" → "סך של 2000 יורו"
    Delivery date: row with "מועד מסירה" → "30.11.2026"
    """
    reg_fee = 0.0
    delivery_date = ""

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            full_row = " ".join(cells)

            if "דמי" in full_row and ("הרשמה" in full_row or "הקמה" in full_row):
                # Extract amount: "סך של 2000 יורו"
                match = re.search(r'סך\s+של\s+([\d,]+)\s*יורו', full_row)
                if match:
                    reg_fee = parse_hebrew_amount(match.group(1))

            if "מועד מסירה" in full_row:
                # Extract date: "30.11.2026"
                match = re.search(r'(\d{1,2}\.\d{1,2}\.\d{4})', full_row)
                if match:
                    delivery_date = match.group(1)

    if reg_fee == 0:
        raise ExtractionError("Cannot extract registration fee from definitions table")

    return reg_fee, delivery_date


class ExtractionError(Exception):
    """Raised when a field cannot be extracted from a document."""
    pass
